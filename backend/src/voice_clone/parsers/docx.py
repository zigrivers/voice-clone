"""DOCX text extraction using python-docx."""

import io
from pathlib import Path

from docx import Document


async def extract_text_from_docx(source: Path | bytes) -> str:
    """Extract text content from a DOCX file.

    Args:
        source: Either a Path to a DOCX file or raw bytes of DOCX content.

    Returns:
        Extracted text content from the document.
    """
    if isinstance(source, bytes):
        source = io.BytesIO(source)

    doc = Document(source)

    text_parts = []

    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n\n".join(text_parts)
